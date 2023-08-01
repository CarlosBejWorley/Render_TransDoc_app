#Importando librerias
import dash
import dash_labs as dl
import dash_bootstrap_components as dbc
from components.navbar.navbar import navbar
from components.form.form import form
from components.doc_alert.doc_alert import alert
import pathlib
#import pdfplumber
import cloudconvert
import requests

import base64
import os
from urllib.parse import quote as urlquote
import glob
from docx import Document
from back_functions.CargosAutoTest_Translate import procesar_doc
from urllib.parse import quote as urlquote


from flask import Flask, render_template, request, redirect, url_for, send_from_directory
from dash import dcc, html
from dash.dependencies import Input, Output, State

#Declarando donde se van a alojar los documentos cargados y descargados
UPLOAD_DIRECTORY = "uploaded_files/"
DOWNLOAD_DIRECTORY = "download_files/"

#Funciones para borrar el directorio de archivos al recargar la pagina
delete_files = glob.glob(UPLOAD_DIRECTORY+'*.docx')
delete_files_pdf = glob.glob(UPLOAD_DIRECTORY+'*.pdf')

for f in delete_files:
    os.remove(f)

for f in delete_files_pdf:
    os.remove(f)

def borrar_downloads():
    delete_downloads = glob.glob(DOWNLOAD_DIRECTORY+'*.docx')
    delete_py = glob.glob(DOWNLOAD_DIRECTORY+'*.py')    
    print("Borrando descargas")
    for f in delete_downloads:
        os.remove(f)
        print("Descarga borrada")
    for f in delete_py:
        os.remove(f)    

borrar_downloads()

# Declarando la instancia de dash 
server = Flask(__name__)
app = dash.Dash(server=server, plugins=[dl.plugins.pages], external_stylesheets=[dbc.themes.BOOTSTRAP], suppress_callback_exceptions=True)
server = app.server

@server.route("/download/<path:path>")
def download(path):
    """Serve a file from the upload directory."""
    return send_from_directory(DOWNLOAD_DIRECTORY, path, as_attachment=True)

#Creando elementos para agregar a la vista principal
docInfo = alert

format_link=  html.A("You can access the available job formats here", id="format_link", target="_blank",href="https://worleyparsons.sharepoint.com/sites/People_grp/GOF/Global%20Job%20Profiles/Forms/AllItems.aspx?viewid=91cf3383%2D27de%2D4bb9%2D9930%2D69f44f6ec8aa")

upload =   dcc.Upload(
            id="upload-data",
            children=html.Div(
                ["Drag and drop or click to select a file to upload."]
            ),
            style={
                "width": "100%",
                "height": "60px",
                "lineHeight": "60px",
                "borderWidth": "1px",
                "borderStyle": "dashed",
                "borderRadius": "5px",
                "textAlign": "center",
                "margin": "10px",
            },
            multiple=True,
        )

button = html.Div(
    [
        dbc.Button(
            "Translate Document", id="translate-button", className="d-grid gap-2 col-6 mx-auto", n_clicks=0
        ),
        
    ]
)

reload_button = html.Div(
    [
        html.A(
            "Translate another document", id="reload_button", className="d-grid gap-2 col-6 mx-auto", n_clicks=0, href="/"
        ),
        
    ]
)

def translate_file(file_path,discipline,education,experience):
        print("Saving File!!!")
        translated_file=procesar_doc(file_path,discipline,education,experience)
        name=file_path.split('/')
        print(name)
        translated_file.save(DOWNLOAD_DIRECTORY+'Translated-'+name[1])

ALLOWED_EXTENSIONS = {'docx', 'pdf', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_file(name, content, discipline, education, experience):
    data = content.encode("utf8").split(b";base64,")[1]
    complete_path = os.path.join(UPLOAD_DIRECTORY, name)
    with open(complete_path, "wb") as fp:
        fp.write(base64.decodebytes(data))
    if complete_path.lower().endswith(".pdf"):
        docx_filename = complete_path.rsplit(".", 1)[0] + ".docx"
        pdf_to_word(complete_path, docx_filename)  # Realiza la conversión de PDF a Word con CloudConvert
        translate_file(docx_filename, discipline, education, experience)
        # os.remove(docx_filename)  # Si deseas eliminar el archivo temporal de Word después de la traducción
    else:
        translate_file(complete_path, discipline, education, experience)


"""Funcion de conversion de pdf a word por medio de pdfPlumber"""
# def pdf_to_word(input_pdf_path, output_word_path):
#     doc = Document()

#     with pdfplumber.open(input_pdf_path) as pdf:
#         # Recorre las páginas del PDF
#         for page in pdf.pages:
#             # Extrae el texto de la página y agrega un párrafo al documento de Word
#             text = page.extract_text()
#             doc.add_paragraph(text)

#             # Extrae las tablas de la página y agrega filas y celdas al documento de Word
#             for table in page.extract_tables():
#                 # Creamos una tabla en el documento
#                 doc_table = doc.add_table(rows=len(table), cols=len(table[0]))

#                 # Agregamos filas y celdas a la tabla
#                 for i, row in enumerate(table):
#                     for j, cell in enumerate(row):
#                         # Verificamos si la celda es nula y la reemplazamos con una cadena vacía
#                         cell_text = cell.strip() if cell else ""
#                         doc_table.cell(i, j).text = cell_text

#     doc.save(output_word_path)

"""Implementacion Cloud Convert"""

API_KEY = "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiODNjODVhZjM5NjBmZGVmYzE2MWQwNDdkMjAxNmRkNjc3MGUxYTIwMDIzMjcwNWM5MTI1ODhiYjBlMmQ4ZTQ4M2EzYWM4MjVjOTg3MzI1YWUiLCJpYXQiOjE2OTA4NjMyNjYuMjkyOTkyLCJuYmYiOjE2OTA4NjMyNjYuMjkyOTkzLCJleHAiOjQ4NDY1MzY4NjYuMjg2NTE1LCJzdWIiOiI2NDU4NDAxMSIsInNjb3BlcyI6WyJ1c2VyLnJlYWQiLCJ1c2VyLndyaXRlIiwidGFzay5yZWFkIiwidGFzay53cml0ZSIsIndlYmhvb2sucmVhZCIsIndlYmhvb2sud3JpdGUiLCJwcmVzZXQucmVhZCIsInByZXNldC53cml0ZSJdfQ.JtTpWz8vuGpRa3U1K2chmK4Y4tsQo3H42SHPF6gW-lJkO5iT_4IQW_ghC5usdjL4Y-NnBw4o_PL3zG_IF1tQtLzQid2llQ0D9wdSFP4nNMR9rzAczMo1JzzZajSPs1t_dAE0V7HesRC5oDBMF1J2v2LU1l79tbxDVY14wKcDOYmht6pnaj-NqK4RjPCwwJ2lK5NEp9h2oT2PKHz7-otxeJWJjk1oGqqRAxF5hGKO_lP51R-SqmCzEI7kJ3u5LOCaJp52PLWk4OdyBO2RgOaQgIyz1JM0CgF7KC9ORnphBYQuloyJGTIb6vZVfjSJBf0oBzN-X8VTivgY9fszgFyhmzgNshys47whuN40Ff-eyRDGz6le2RJkMEiGvjc6UZpVSAgjFZIP8we2PrZ0MhviBFzUYhMJ3WPGkmm-BlCpyrshyg605T7KtvB0zPe1WZtRO6npDo4f7XMjHa0AYfwDAhX7wUA4zJ48D_XuWJZsuQ7LZWQCmtXYTvJhsjRZPwtE0OQl91i01wE7YO2xk6UYDXdo1RN3O6JIENYpmQld2LMVlpjbBG1UTvPvQV9CX8JUI5CHCQKaODDqa1GWAcKdNzKSvcFqIzde3oAD6MxKPKbRnh9PrsUePMHuY26Zr9pVJNZL9RBhnRccTlJFtXgUSLlv9ZFeoIYPIOpUpkXtAG4"
#API_KEY = "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiYzNiOTZkYzg5OTE2ZjhhZDM3MTJhMmUxOGEzYzhlMzgxYjBiZGIxNjM5MGJkYTc5OWE3MjIzMGVkYzYwNGY0ZTgzN2I1YTRkNWM3YmYyMTkiLCJpYXQiOjE2OTA4MzQ1NDYuMjAxODA4LCJuYmYiOjE2OTA4MzQ1NDYuMjAxODEsImV4cCI6NDg0NjUwODE0Ni4xOTQwMzcsInN1YiI6IjY0NTUzODE3Iiwic2NvcGVzIjpbIndlYmhvb2sud3JpdGUiLCJ3ZWJob29rLnJlYWQiLCJ0YXNrLndyaXRlIiwidGFzay5yZWFkIiwidXNlci53cml0ZSIsInVzZXIucmVhZCIsInByZXNldC5yZWFkIiwicHJlc2V0LndyaXRlIl19.MN_kyRQn21KIbZ0W0DGqB6J_m1Rri2P5oD-uZeaXue-TGLNbByYBiik9yissmZwv_fuw1NuA68TtMZqZM-dbwVz5pdLrmDkTK5iX50kijt3C5Y2QAj5GH6MlDGkqJAaZ0bCNO_8sEATBwHbQDuhVAL01f1EC3ei6_wq_EgxF8CLlQxbtY5WhwLUjJLF50YKVW6BaakPT2jbd7RVLpX1CTT2mftLtK7KUn5chgJySjDygNCn9QKsCBmpHVZg5jK7gZRCYElHg6O3Ad_OM7Aq_vPkw4JS3aqqY4wvw427jaXd13dWW4FonwZZhZgACdCLewQPNbWtUgNnYdd9KvFFnnPPCX_f9q1VP80IRpqrk8bWihczv9aBFZFxwt8lRHV9zVubKR5xKUCieC63e8KD5ivsiym9UIYSzJGrCltMwYN9A80FJU2fi2Jt66amhOG_0F5L43L91lh_xoRp8aDYDZ7tTDIUEuc--zg531TuKCD3tZKNJ0YWq4HsSMqUt_Nu9NrQPyYaL951RPEQ7e5pRc8ziKQwKyI6k2iFnlm7bpXwUfZ4fpdfH17gTa-fFzpXyKDPoQDrJgB1sk2gMkSbis7cgFmav2NyZKuEkGyGvKrD5ejChn4qbdrlmOKq-_wcRfClBUSTt-ZX6ULjLeNyrxDIbjKb_9TC9lY62zhSVEiI"
cloudconvert.configure(api_key=API_KEY)

def pdf_to_word(input_pdf_path, output_word_path):
    # Crea el trabajo para importar el archivo PDF y convertirlo a Word
    job = cloudconvert.Job.create(payload={
        "tasks": {
            "import-my-file": {
                "operation": "import/upload"
            },
            "convert-my-file": {
                "operation": "convert",
                "input": "import-my-file",
                "output_format": "docx"  # Especifica el formato de salida como docx para Word
            },
            "export-my-file": {
                "operation": "export/url",
                "input": "convert-my-file"
            }
        }
    })

    # Crea una nueva tarea de importación para cargar el archivo PDF
    import_task = cloudconvert.Task.upload(file_name=input_pdf_path, task=job['tasks'][0])

    # Espera a que se complete el trabajo completo (importación y conversión)
    job = cloudconvert.Job.wait(id=job['id'])

    # Verifica si el trabajo fue exitoso
    if job['status'] == 'finished':
        # Obtiene la URL de descarga del archivo Word convertido
        #print(job['tasks'][0]['id'])
        #export_task_id = job['tasks']['export-my-file']['id']
        export_task_id = job['tasks'][0]['id']
        export_task = cloudconvert.Task.find(id=export_task_id)
        export_url = export_task['result']['files'][0]['url']

        response = requests.get(export_url)

        # Verifica si la solicitud fue exitosa
        if response.status_code == 200:
            # Escribe el contenido binario en el archivo
            with open(output_word_path, 'wb') as output_file:
                output_file.write(response.content)
            print("PDF converted to Word successfully!")
        else:
            print("Error downloading the converted file.")



def uploaded_files():
    """List the files in the download directory."""
    files = []
    for filename in os.listdir(DOWNLOAD_DIRECTORY):
        path = os.path.join(DOWNLOAD_DIRECTORY, filename)
        if os.path.isfile(path):
            files.append(filename)
    return files

def file_download_link(filename):
    """Create a Plotly Dash 'A' element that downloads a file from the app."""
    location = "/download/{}".format(urlquote(filename))
    return html.A(filename, href=location)

def allowed_file_error_message():
    return html.Div(
        "Invalid file type. Allowed file types are: " + ", ".join(ALLOWED_EXTENSIONS),
        style={
                "color": "#A83B42",
                "background-color" : "#ea868f",
                "border-radius" : ".25rem",
                "padding" : "1rem 1rem",
                "border" : "solid 1px transparent",
                "border-color" : "#F35C75"
            },
        className="text-center"
    )   

#Vista Principal
app.layout = dbc.Container(
    [   navbar,
        dbc.Container([
            dbc.Row(html.H3("Welcome to JDT Bot – Job Description Translator Bot", className="text-center",)),
            dbc.Row(html.H4("Follow these steps:")),
            dbc.Row(html.H5("1. Upload your document to be translated and exported to new format: ")),
            dbc.Row(format_link, style={"textAlign": "center"}),
            dbc.Row(upload),
            dbc.Row(allowed_file_error_message(), id="allowed-file-message"),
            dbc.Row(docInfo),
            dbc.Row(html.H5("2. Define extra content to be added:")),
            dbc.Row(form),
            dbc.Row(html.H5("3. Translate your document:")),
            dbc.Row(button),
            dbc.Row(html.H5("4. Download your Document:")),
            dbc.Row(html.H5(id="download_File", style={"textAlign": "center"})),
            dbc.Row(dbc.Collapse(html.H5(reload_button),id="reload_collapse",is_open=False), className="text-center")
        ]            
        )      
        
    ],
    className="dbc",
    fluid=True,
)

""" Desplegar alerta al subir un documento """
@app.callback(
    [Output("doc_alert", "children"), Output("doc_alert", "is_open"), Output("allowed-file-message", "children")],
    [Input("upload-data", "contents")],
    [State("upload-data", "filename"), State("doc_alert", "is_open")]
)
def update_alert(contents, uploaded_filenames, is_open):
    if not contents:
        return "Upload a document to translate", False, None

    name = uploaded_filenames[0] if uploaded_filenames else None
    content = contents[0] if contents else None

    if not allowed_file(name):
        return "Upload a document to translate", False, allowed_file_error_message()

    if name and is_open is not True:
        return "File: " + name + " processed successfully!", True, None

    return "Upload a document to translate", False, None




"""Guardar el documento subido, traducirlo y generar un nuevo documento descargable"""

@app.callback(
    Output("download_File", "children"),
    [Input("translate-button", "n_clicks")],
    [State("upload-data", "filename"),State("upload-data", "contents"),State("discipline-row", "value"),State("education-radios", "value"),State("experience-radios", "value")]
    #,Input("upload-data", "filename"), Input("upload-data", "contents")
)
def update_output(click,uploaded_filenames, uploaded_file_contents,discipline,education,experience):
    
    if click is not None:

        if uploaded_filenames is not None and uploaded_file_contents is not None:
            for name, data in zip(uploaded_filenames, uploaded_file_contents):
                save_file(name, data, discipline, education, experience)

        files = uploaded_files()
        if len(files) == 0 :
            return "Load a file on first step to translate"
        else:
            return [file_download_link(files[0])]
            
"""Desplegar opción de nuevo documento"""
@app.callback(
    Output("reload_collapse","is_open"),
    Input("download_File","children")
)
def update_reload(content):
    if content == "Load a file on first step to translate":
        return False
    else:
        return True

"""Opción para traducir un nuevo documento"""
@app.callback(
    Output("reload_button", "children"),
    Input("reload_button","n_clicks")
)
def reset_upload(click):
    if click is not None:
        borrar_downloads()
        return "Translate another document"
    

# Declarando el servidor
if __name__ == "__main__":
    app.run_server(debug=True, host='127.0.0.1')