#%%
#Paquetes
#!pip install python-docx
#!pip install google_trans_new
import pandas as pd
import numpy as np
from docx import Document
from google_trans_new import google_translator  


# %%
plantilla = Document('HR-FRM-0001 FORMATO DESCRIPTIVO DE CARGO_Esp (Plantilla).docx')
#docinfo = Document('Cargos/5A - Senior Engineer.docx')

#Iniciando variable para usar traductor de google
translator = google_translator()  


#%%
#Función para cambiar texto en el documento (fuera de tablas)
def cambiaTexto(llaveOriginal,textoNuevo,doc):
    for paragraph in doc.paragraphs:
        if llaveOriginal in paragraph.text:
            #print(paragraph.text)
            paragraph.text = textoNuevo


#Función para cambiar texto dentro de tablas
def cambiaTextoTabla(llaveOriginal,textoNuevo,doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if llaveOriginal in paragraph.text:
                        spanishText = traducirTexto(textoNuevo)
                        paragraph.text = paragraph.text.replace(llaveOriginal, spanishText)


#Función para traducir texto desde inglés a español
def traducirTexto(textOriginal):
    spaText=translator.translate(textOriginal, lang_src='en', lang_tgt='es')
    return spaText


def procesar_doc(path):


    docinfo = Document("Cargos/"+path)
    #%%
    #Obtener información de la tabla del documento original
    info = [''] #Se guardan en esta lista

    for table in docinfo.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    try:
                        if info[len(info)-1] != str(paragraph.text):
                            info.append(str(paragraph.text))
                            #print(paragraph.text)
                    except:
                        pass


#info


    #%%
    #Los títulos en el documento original y sus respectivos tags en la plantilla
    fields = {'Role':'&GLOBAL.JOB.TITLE&',
        'Job Family':'&JOB.FAMILY&',
        'Sub-Family':'&SUB.FAMILY&',
        'General Purpose of Role':'&PROPOSITO.GENERAL&',
        'Typically Reports To':'&REPORTA.A&',
        'Example Global Titles':'',
        'Alternative Location Titles':'',
        'Business Development Support':'',
        'Business Operations':'',
        'Technical Competency':'',
        'Project Execution':'',
        'Project Deliverables':'',
        'Procurement Support':'',
        'Site Support':'',
        'Global Level Summary':'&OBJETIVO.GENERAL&',
        'Qualifications, Accreditation, Training (Essential)':'',
        'Qualifications, Accreditation, Training (Desirable)':'',
        'Job Specific Knowledge / Experience (Essential)':'&PRINCIPALES.RESPONSABILIDADES&',
        'Job Specific Knowledge / Experience (Desirable)':'',
        'Decision Making':'&TOMA.DE.DECISIONES&',
        'Supervision Received':'&NIVEL.SUPERVISION&',
        'Supervision Authority':'&AUTORIDAD.SUPERVISION&',
        'Communication':'&COMUNICACION&',
        'Systems & Tools (Essential)':'&SPH.ESENCIALES&',
        'Systems & Tools (Desirable)':'&SPH.DESEABLES&',
        'HSE Capability':'',
        'People Skills':'&HABILIDADES.PERSONALES&'}

    #%%
    #Crear un DataFrame que contenga los títulos, los tags, y más adelante el texto
    campos = pd.DataFrame.from_dict(fields, orient='index') #Crear el DataFrame inicial
    campos.reset_index(inplace=True)
    campos.columns = ['Original','TagNuevo']
    campos['Texto'] = ''

    

    #Incluir el texto en la plantilla
    for i in range(0,len(info)-1):
        if info[i] in fields:
            textoEnPos = campos[campos['Original']==info[i]]['Texto']
            if textoEnPos.iloc[0] != '':    #Si ya hay texto, agregar un salto de línea
                textoNuevo = textoEnPos + '\n' + info[i+1]
            else:
                textoNuevo = info[i+1]
            campos.loc[campos['Original']==info[i],'Texto'] = textoNuevo
            #print(str(i) + '. ' + info[i] + ' == ' + info[i+1])


    #campos['Texto'] = campos['Texto'].str.capitalize()
    campos



    #%%
    #Cargar texto a la plantilla (cargado en Python), reemplazando los tags
    for tag in campos['TagNuevo']:
        if tag != '':
            cambiaTextoTabla(tag,campos.loc[campos['TagNuevo']==tag,'Texto'].iloc[0],plantilla)



    # %%
    #Guardar nuevo archivo
    plantilla.save('Traducidos/'+path)
    # %%
