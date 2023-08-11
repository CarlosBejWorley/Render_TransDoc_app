#from matplotlib.pyplot import text
import pandas as pd
import numpy as np
from docx import Document
from deep_translator import GoogleTranslator
import pathlib
#import pdfplumber
import cloudconvert
import requests
import hashlib
import os

'''
El funcionamiento del algoritmo de traducción es el siguiente:
    1. Se carga la plantilla que funcionara como borrador del documento traducido en el nuevo formato
    2. Buscar los textos en el docuemnto original (Esten dentro o fuera de una tabla)
    3. Traducir los textos pertinentes a español
    4. Buscar la posición dentro del borrador donde debe estar el texto
    5. Colocar cada texto traducido en el borrador, colocar las selecciones de los campos de disicplina , educación y experiencia seleccionadas por el usuario
    6. Se retorna el documento en el nuevo formato con el nombre del archivo original
'''


# Cargando el documento que funcionara como plantilla
plantilla = Document('plantillas/Plantilla.docx')  #Tiene que hacer esto con cada nuevo documento

#Iniciando variable para usar traductor de google
translator = GoogleTranslator(source="english",target="es")

#Función para cambiar texto por fuera de las tablas del documento original al nuevo formato 
def cambiaTexto(llaveOriginal,textoNuevo,doc):
    for paragraph in doc.paragraphs:
        if llaveOriginal in paragraph.text:
            #print(paragraph.text)
            paragraph.text = textoNuevo

#Función para cambiar texto de las tablas del documento original al nuevo formato 
def cambiaTextoTabla(llaveOriginal,textoNuevo,doc,disciplina,educacion,experiencia): 

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if llaveOriginal in paragraph.text:
                        spanishText = traducirTexto(textoNuevo)
                        #spanishText = textoNuevo
                        paragraph.text = paragraph.text.replace(llaveOriginal, spanishText)

                    #Buscando y reemplazando los campos de disciplina, educacion y experiencia
                    if "&DISCIPLINA&" in paragraph.text:
                        paragraph.text = paragraph.text.replace(paragraph.text,translator.translate(disciplina))    

                    if "&EDU.{}&".format(educacion) in paragraph.text:                        
                        paragraph.text = paragraph.text.replace(paragraph.text,"X")
                    elif "&EDU." in paragraph.text:
                        paragraph.text = paragraph.text.replace(paragraph.text,"")
                        
                    if "&EXP.{}&".format(experiencia) in paragraph.text:
                        paragraph.text = paragraph.text.replace(paragraph.text,"X")                        
                    elif "&EXP." in paragraph.text:
                        paragraph.text = paragraph.text.replace(paragraph.text,"")

#Función para traducir texto desde el idioma origen a español
def traducirTexto(textOriginal):
    spaText=translator.translate(textOriginal)
    return spaText    

#Función para pasar el texto en los campos del documento original a la plantilla
def procesar_doc(path,disciplina='Engineering',educacion='1',experiencia='3'):
    
    docinfo = Document(path)

    #Obtener información de la tabla del documento original
    info = [''] #Se guardan en esta lista

    for table in docinfo.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    try:
                        if info[len(info)-1] != str(paragraph.text):
                            info.append(str(paragraph.text))
                            #print(paragraph.text)
                    except:
                        pass

    info

    #
    #Los títulos en el documento original y sus respectivos tags en la plantilla
    fields = {'Role':'&CARGO.LOCAL&',
        'Job Family':'&JOB.FAMILY&',
        'Sub-Family':'&SUB.FAMILY&',
        'Global Job Title': '&GLOBAL.JOB.TITLE&', 
        'General Purpose of Role':'&PROPOSITO.GENERAL&',
        'Typically Reports To':'&REPORTA.A&',
        'Example Global Titles':'',
        'Alternative Location Titles':'',
        'Business Development Support':'',
        'Business Operations':'',
        'Technical Competency':'',
        'Project Execution':'',
        'Project Deliverables':'',
        'Procurement Support':'',
        'Site Support':'',
        'Global Level Summary':'&OBJETIVO.GENERAL&',
        'Qualifications, Accreditation, Training (Essential)':'',
        'Qualifications, Accreditation, Training (Desirable)':'',
        'Job Specific Knowledge / Experience (Essential)':'&PRINCIPALES.RESPONSABILIDADES&',
        'Job Specific Knowledge / Experience (Desirable)':'',
        'Decision Making':'&TOMA.DE.DECISIONES&',
        'Supervision Received':'&NIVEL.SUPERVISION&',
        'Supervision Authority':'&AUTORIDAD.SUPERVISION&',
        'Communication':'&COMUNICACION&',
        'Systems & Tools (Essential)':'&SPH.ESENCIALES&',
        'Systems & Tools (Desirable)':'&SPH.DESEABLES&',
        'HSE Capability':'',
        'People Skills':'&HABILIDADES.PERSONALES&',
        
        }

    #
    #Crear un DataFrame que contenga los títulos, los tags, y más adelante el texto
    campos = pd.DataFrame.from_dict(fields, orient='index') #Crear el DataFrame inicial
    campos.reset_index(inplace=True)
    campos.columns = ['Original','TagNuevo']
    campos['Texto'] = ''


    #Incluir el texto en la plantilla
    for i in range(0,len(info)-1):
        if info[i] in fields:
            textoEnPos = campos[campos['Original']==info[i]]['Texto']
            if textoEnPos.iloc[0] != '':    #Si ya hay texto, agregar un salto de línea
                textoNuevo = textoEnPos + '\n' + info[i+1]
            else:
                textoNuevo = info[i+1]
            campos.loc[campos['Original']==info[i],'Texto'] = textoNuevo
            #print(str(i) + '. ' + info[i] + ' == ' + info[i+1])


    #campos['Texto'] = campos['Texto'].str.capitalize()
    campos


    #Cargar texto a la plantilla (cargado en Python), reemplazando los tags
    for tag in campos['TagNuevo']:
        if tag != '':
            cambiaTextoTabla(tag,campos.loc[campos['TagNuevo']==tag,'Texto'].iloc[0],plantilla,disciplina,educacion,experiencia)


    #Retornar archivo traducido
    return plantilla
    #plantilla.save('downloads_files/'+path)

api_keys = [
    "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiODNjODVhZjM5NjBmZGVmYzE2MWQwNDdkMjAxNmRkNjc3MGUxYTIwMDIzMjcwNWM5MTI1ODhiYjBlMmQ4ZTQ4M2EzYWM4MjVjOTg3MzI1YWUiLCJpYXQiOjE2OTA4NjMyNjYuMjkyOTkyLCJuYmYiOjE2OTA4NjMyNjYuMjkyOTkzLCJleHAiOjQ4NDY1MzY4NjYuMjg2NTE1LCJzdWIiOiI2NDU4NDAxMSIsInNjb3BlcyI6WyJ1c2VyLnJlYWQiLCJ1c2VyLndyaXRlIiwidGFzay5yZWFkIiwidGFzay53cml0ZSIsIndlYmhvb2sucmVhZCIsIndlYmhvb2sud3JpdGUiLCJwcmVzZXQucmVhZCIsInByZXNldC53cml0ZSJdfQ.JtTpWz8vuGpRa3U1K2chmK4Y4tsQo3H42SHPF6gW-lJkO5iT_4IQW_ghC5usdjL4Y-NnBw4o_PL3zG_IF1tQtLzQid2llQ0D9wdSFP4nNMR9rzAczMo1JzzZajSPs1t_dAE0V7HesRC5oDBMF1J2v2LU1l79tbxDVY14wKcDOYmht6pnaj-NqK4RjPCwwJ2lK5NEp9h2oT2PKHz7-otxeJWJjk1oGqqRAxF5hGKO_lP51R-SqmCzEI7kJ3u5LOCaJp52PLWk4OdyBO2RgOaQgIyz1JM0CgF7KC9ORnphBYQuloyJGTIb6vZVfjSJBf0oBzN-X8VTivgY9fszgFyhmzgNshys47whuN40Ff-eyRDGz6le2RJkMEiGvjc6UZpVSAgjFZIP8we2PrZ0MhviBFzUYhMJ3WPGkmm-BlCpyrshyg605T7KtvB0zPe1WZtRO6npDo4f7XMjHa0AYfwDAhX7wUA4zJ48D_XuWJZsuQ7LZWQCmtXYTvJhsjRZPwtE0OQl91i01wE7YO2xk6UYDXdo1RN3O6JIENYpmQld2LMVlpjbBG1UTvPvQV9CX8JUI5CHCQKaODDqa1GWAcKdNzKSvcFqIzde3oAD6MxKPKbRnh9PrsUePMHuY26Zr9pVJNZL9RBhnRccTlJFtXgUSLlv9ZFeoIYPIOpUpkXtAG4",
    "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiYzNiOTZkYzg5OTE2ZjhhZDM3MTJhMmUxOGEzYzhlMzgxYjBiZGIxNjM5MGJkYTc5OWE3MjIzMGVkYzYwNGY0ZTgzN2I1YTRkNWM3YmYyMTkiLCJpYXQiOjE2OTA4MzQ1NDYuMjAxODA4LCJuYmYiOjE2OTA4MzQ1NDYuMjAxODEsImV4cCI6NDg0NjUwODE0Ni4xOTQwMzcsInN1YiI6IjY0NTUzODE3Iiwic2NvcGVzIjpbIndlYmhvb2sud3JpdGUiLCJ3ZWJob29rLnJlYWQiLCJ0YXNrLndyaXRlIiwidGFzay5yZWFkIiwidXNlci53cml0ZSIsInVzZXIucmVhZCIsInByZXNldC5yZWFkIiwicHJlc2V0LndyaXRlIl19.MN_kyRQn21KIbZ0W0DGqB6J_m1Rri2P5oD-uZeaXue-TGLNbByYBiik9yissmZwv_fuw1NuA68TtMZqZM-dbwVz5pdLrmDkTK5iX50kijt3C5Y2QAj5GH6MlDGkqJAaZ0bCNO_8sEATBwHbQDuhVAL01f1EC3ei6_wq_EgxF8CLlQxbtY5WhwLUjJLF50YKVW6BaakPT2jbd7RVLpX1CTT2mftLtK7KUn5chgJySjDygNCn9QKsCBmpHVZg5jK7gZRCYElHg6O3Ad_OM7Aq_vPkw4JS3aqqY4wvw427jaXd13dWW4FonwZZhZgACdCLewQPNbWtUgNnYdd9KvFFnnPPCX_f9q1VP80IRpqrk8bWihczv9aBFZFxwt8lRHV9zVubKR5xKUCieC63e8KD5ivsiym9UIYSzJGrCltMwYN9A80FJU2fi2Jt66amhOG_0F5L43L91lh_xoRp8aDYDZ7tTDIUEuc--zg531TuKCD3tZKNJ0YWq4HsSMqUt_Nu9NrQPyYaL951RPEQ7e5pRc8ziKQwKyI6k2iFnlm7bpXwUfZ4fpdfH17gTa-fFzpXyKDPoQDrJgB1sk2gMkSbis7cgFmav2NyZKuEkGyGvKrD5ejChn4qbdrlmOKq-_wcRfClBUSTt-ZX6ULjLeNyrxDIbjKb_9TC9lY62zhSVEiI",
    "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiOWEwZDFiYTlmYzc3MDdjMThkYTVlNWQwODM3ZGM5OTYzZDU5YTE1NzBiYzI1NDk2MmMwMWQzOTY1MzM5MjIwMTYwNTBhNWNkZWQ0MDgwZTEiLCJpYXQiOjE2OTE2NzY0MzkuMTkyNDY2LCJuYmYiOjE2OTE2NzY0MzkuMTkyNDY4LCJleHAiOjQ4NDczNTAwMzkuMTg3MDMxLCJzdWIiOiI2NDU4OTYyOCIsInNjb3BlcyI6WyJ1c2VyLnJlYWQiLCJ1c2VyLndyaXRlIiwidGFzay5yZWFkIiwidGFzay53cml0ZSIsIndlYmhvb2sucmVhZCIsIndlYmhvb2sud3JpdGUiLCJwcmVzZXQucmVhZCIsInByZXNldC53cml0ZSJdfQ.qsEXfOxZtq5tfqKOShLhc9JGQRhXfA9MWK6ahFV1WpZaTjSq4WICdg89-xBFasDOKJtVuhfgfvzfH8O0HjLYacyg4ky7qhnSzVlHLaZU887Dj_-T96m8uagTZssM-BK2eVS4jvUcJdtB0YrYiVbOSV6N3ZIkCagYF45_1ZxuHfi31DwB7jpDzkVWRtl3tvVHR3VUAdmIiP5PuHxcxSAmXiOoGdgFdNIuFrMVKjGwyEyqz2yWWniIj8jQECsqLC7LMksStgDTmTRSvxBfaFv_9jj0MKsajs3u2RfYkt9gKk0ufnXvjPjjwKFAxHgl0dZ_Zd2cjJhoH1fD5S7HfMEeDzAj-45_QMPJhQ9eHdDDwBfw3LFjBhaKeH1ylNN8X1hfNAWuN9JS13qOrtTqqLKI7cjxTLRTa9O0wVd98x6h5SCFifxUy54eRYiFYG6EKeXvFBA5TorrkrvvZd4fY9Z9VlPV00xzKfRm8fXDW1rSOMMCXQf5qGjnSRs9UVbhF9dZkZpYSi4rzIlP11tYsNlW1DaCO8PWbgp4Pq4IInHovoJRitEDO93lKhjELIX4zkWHlEe-JqKtPP-Kmr58TJhqcVWPUzz23HSZjYCKpmolE596mmItO8BrfdqGFYWtkyX-pcWIpJhIj91iN9iOjmHBXdXABPgCe4Ega8fr17Z08oM",
    "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiZDU5NmZiOGRkOGJjNDdjMTk0NzFhNWYzMzVjNTgxZjU5ZGQ0ZDUyMjQ0NDJmYzI5OTc3ZTk3MzQyMTE1MWJhYzE4MTZlOGZkNzI5ZjYyZjQiLCJpYXQiOjE2OTE2ODYyNTYuNjkzMzk0LCJuYmYiOjE2OTE2ODYyNTYuNjkzMzk2LCJleHAiOjQ4NDczNTk4NTYuNjg4Njg5LCJzdWIiOiI2NDcxNDIwNyIsInNjb3BlcyI6WyJ1c2VyLndyaXRlIiwidXNlci5yZWFkIiwidGFzay5yZWFkIiwidGFzay53cml0ZSIsIndlYmhvb2sucmVhZCIsIndlYmhvb2sud3JpdGUiLCJwcmVzZXQucmVhZCIsInByZXNldC53cml0ZSJdfQ.EUMso5M0HAI38X32UvHSIpHz5oGY0-121Hvk8Ve3FPyHUVw6bxV2NOt-dMMMcR9fyDAlFgnmUCTnN8-SQTAUFdlud3C-tFYZAsvyptfY-FN5jd4Tz8w6n0E3aUGsGjVaQ7PVDJ3WIuzh1ElRow_RWJxaQRN80yv2_28hXQuhBivOC6e3QRD2olRmkCwq8Sru9dQ57SOrzM7RY9Nhd7jRf53mkzzYUxFVP-ypPDcUPD3-ZnK7HT58K8W-oaiJd85FA0x1X77WRCb--KnOe96VhtIey_HHHY0gWht09mIdqyOLa-5gzi_se4uiKsozIKuS6R-epAm03gh7Dt1TuEubMbIPpvvk1vMGgB0yEyfQiMv2j-x4ac0L1UWn1P2vjsuph8d9yzMKEaRsvYHtocWZB1Tr1nVarakD12bSSZv_XF-bVlBWVPYxcPAwBwmLBbU6jb3MFB8JKFurYltDuPi7a6RrFZaNWwDxSWm9I0t7kU05YJQf4iS8q-lWCf0MCPO-r0gZPQ-z4WYglw0BE6pYym2k_FG_xsFL9XNxK9MNROWGZcDZNX57G-MgXA1IK4WCMNxtUbefa8009XqehiOT-mg_e-RlHM5b8Kh5A9ztpY3mxOypvfNfrM0kMAh7fqkBQr9CaIkxNiuVq6ozGLt2JvSERDq1cqUoIGdc1yr7DYQ",
    "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiZDUwZDY2MjE2MTBiMDY3ZmQwYjAxNzNmYmUyMGU5YzYyZjA5YzIxNGRhMTVjNjBhZDUzMGNlZWNmNWU0MDRjNjcyZjQxNDAxNjYwZDQ1ZTYiLCJpYXQiOjE2OTE2ODYyMTkuMDEzMjYsIm5iZiI6MTY5MTY4NjIxOS4wMTMyNjEsImV4cCI6NDg0NzM1OTgxOS4wMDgxODgsInN1YiI6IjY0NzE0MTk3Iiwic2NvcGVzIjpbInVzZXIucmVhZCIsInVzZXIud3JpdGUiLCJ0YXNrLnJlYWQiLCJ0YXNrLndyaXRlIiwid2ViaG9vay5yZWFkIiwid2ViaG9vay53cml0ZSIsInByZXNldC5yZWFkIiwicHJlc2V0LndyaXRlIl19.Lasoj_MiDWMOxjoAhIbdb2JpqaZ43eAw0cpc3UDTOzjtlP1iEHP9wxE9-kZjbcSDtc3pOebVMVbKnymVaa8jWZXk7P2crNZCWMCeLNcm96yWvAIUut6eQqgtzC9AuSzSJ31vAVbDlMrg3ljNTk05d1lED7R6h8TeqYxUQZDmZHFgp7A4mVhKkBOEi_gnr80YZCiXeD6qKWiTRG-B-zcd_wO2oNJASItUA3tLrdK8SwDXTQz2F106zIwh4UjRfNYSpplKBCX0_kKLWBGi9qXsrKgOHzEC-Y-scKLXUALQN6YHsoL8kK8RZModGeMmH_zpXrKgtG5qZfQZEwkN6ie4SnDV3bbfAjCgcAZ_eWWvSITM8-R3xax-yNm5np1Qaflh6Pq_hFvXuhjvQu_Hz1uJ6Incnz34IVC4iElLbTPf1KB90OiRHgFllTSYl5RnD0jRYs_w5xaEKiAW7fqfaHul3_NXS0VdIYiVNTuMRwaF66DhLKo66yeRblEiWHIoZ9V377zjA92MXeyce919CZrb66oCrbXwyMjHl9TC4elvKyVpupjnP6zKseSmi548VDw9nv6JSIEzw1wQJQzpjFqFva-XcBVXFKn8gUrPbss4XBQ8VCgiD5lpjwH2FfiVtgNdjUS0gstJBk4Jj4lv78nJwiZ0UZnvZbMNyKLHNUprRh0"
]

def hash_api_key(api_key, salt=None, iterations=100000):
    if salt is None:
        salt = os.urandom(16)  # Generates a random 16-byte salt

    hashed_api_key = hashlib.pbkdf2_hmac('sha256', api_key.encode('utf-8'), salt, iterations)

    return salt, iterations, hashed_api_key

def deshash_api_key(api_key, salt, iterations):
    hashed_api_key_to_verify = hashlib.pbkdf2_hmac('sha256', api_key.encode('utf-8'), salt, iterations)
    return hashed_api_key_to_verify

"""Implementacion Cloud Convert"""

#API_KEY = "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiODNjODVhZjM5NjBmZGVmYzE2MWQwNDdkMjAxNmRkNjc3MGUxYTIwMDIzMjcwNWM5MTI1ODhiYjBlMmQ4ZTQ4M2EzYWM4MjVjOTg3MzI1YWUiLCJpYXQiOjE2OTA4NjMyNjYuMjkyOTkyLCJuYmYiOjE2OTA4NjMyNjYuMjkyOTkzLCJleHAiOjQ4NDY1MzY4NjYuMjg2NTE1LCJzdWIiOiI2NDU4NDAxMSIsInNjb3BlcyI6WyJ1c2VyLnJlYWQiLCJ1c2VyLndyaXRlIiwidGFzay5yZWFkIiwidGFzay53cml0ZSIsIndlYmhvb2sucmVhZCIsIndlYmhvb2sud3JpdGUiLCJwcmVzZXQucmVhZCIsInByZXNldC53cml0ZSJdfQ.JtTpWz8vuGpRa3U1K2chmK4Y4tsQo3H42SHPF6gW-lJkO5iT_4IQW_ghC5usdjL4Y-NnBw4o_PL3zG_IF1tQtLzQid2llQ0D9wdSFP4nNMR9rzAczMo1JzzZajSPs1t_dAE0V7HesRC5oDBMF1J2v2LU1l79tbxDVY14wKcDOYmht6pnaj-NqK4RjPCwwJ2lK5NEp9h2oT2PKHz7-otxeJWJjk1oGqqRAxF5hGKO_lP51R-SqmCzEI7kJ3u5LOCaJp52PLWk4OdyBO2RgOaQgIyz1JM0CgF7KC9ORnphBYQuloyJGTIb6vZVfjSJBf0oBzN-X8VTivgY9fszgFyhmzgNshys47whuN40Ff-eyRDGz6le2RJkMEiGvjc6UZpVSAgjFZIP8we2PrZ0MhviBFzUYhMJ3WPGkmm-BlCpyrshyg605T7KtvB0zPe1WZtRO6npDo4f7XMjHa0AYfwDAhX7wUA4zJ48D_XuWJZsuQ7LZWQCmtXYTvJhsjRZPwtE0OQl91i01wE7YO2xk6UYDXdo1RN3O6JIENYpmQld2LMVlpjbBG1UTvPvQV9CX8JUI5CHCQKaODDqa1GWAcKdNzKSvcFqIzde3oAD6MxKPKbRnh9PrsUePMHuY26Zr9pVJNZL9RBhnRccTlJFtXgUSLlv9ZFeoIYPIOpUpkXtAG4"
#API_KEY = "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiYzNiOTZkYzg5OTE2ZjhhZDM3MTJhMmUxOGEzYzhlMzgxYjBiZGIxNjM5MGJkYTc5OWE3MjIzMGVkYzYwNGY0ZTgzN2I1YTRkNWM3YmYyMTkiLCJpYXQiOjE2OTA4MzQ1NDYuMjAxODA4LCJuYmYiOjE2OTA4MzQ1NDYuMjAxODEsImV4cCI6NDg0NjUwODE0Ni4xOTQwMzcsInN1YiI6IjY0NTUzODE3Iiwic2NvcGVzIjpbIndlYmhvb2sud3JpdGUiLCJ3ZWJob29rLnJlYWQiLCJ0YXNrLndyaXRlIiwidGFzay5yZWFkIiwidXNlci53cml0ZSIsInVzZXIucmVhZCIsInByZXNldC5yZWFkIiwicHJlc2V0LndyaXRlIl19.MN_kyRQn21KIbZ0W0DGqB6J_m1Rri2P5oD-uZeaXue-TGLNbByYBiik9yissmZwv_fuw1NuA68TtMZqZM-dbwVz5pdLrmDkTK5iX50kijt3C5Y2QAj5GH6MlDGkqJAaZ0bCNO_8sEATBwHbQDuhVAL01f1EC3ei6_wq_EgxF8CLlQxbtY5WhwLUjJLF50YKVW6BaakPT2jbd7RVLpX1CTT2mftLtK7KUn5chgJySjDygNCn9QKsCBmpHVZg5jK7gZRCYElHg6O3Ad_OM7Aq_vPkw4JS3aqqY4wvw427jaXd13dWW4FonwZZhZgACdCLewQPNbWtUgNnYdd9KvFFnnPPCX_f9q1VP80IRpqrk8bWihczv9aBFZFxwt8lRHV9zVubKR5xKUCieC63e8KD5ivsiym9UIYSzJGrCltMwYN9A80FJU2fi2Jt66amhOG_0F5L43L91lh_xoRp8aDYDZ7tTDIUEuc--zg531TuKCD3tZKNJ0YWq4HsSMqUt_Nu9NrQPyYaL951RPEQ7e5pRc8ziKQwKyI6k2iFnlm7bpXwUfZ4fpdfH17gTa-fFzpXyKDPoQDrJgB1sk2gMkSbis7cgFmav2NyZKuEkGyGvKrD5ejChn4qbdrlmOKq-_wcRfClBUSTt-ZX6ULjLeNyrxDIbjKb_9TC9lY62zhSVEiI"

def pdf_to_word(input_pdf_path, output_word_path, api_keys):
    for api_key in api_keys:
        try:

            #salt, iterations, hashed_api_key = hash_api_key(api_key)
            #stored_api_key_info = (hashed_api_key, salt, iterations)
            #print ("Hash" , stored_api_key_info)

            #api = deshash_api_key (*stored_api_key_info)
            #print ("Deshash", api_key)
            #Configura la api key para ser usada en el cloudconvert
            cloudconvert.configure(api_key= api_key)

            #Crea un job en cloud convert por medio de un payload
            job = cloudconvert.Job.create(payload={
                "tasks": {
                    "import-my-file": {
                        "operation": "import/upload"
                    },
                    "convert-my-file": {
                        "operation": "convert",
                        "input": "import-my-file",
                        "output_format": "docx"
                    },
                    "export-my-file": {
                        "operation": "export/url",
                        "input": "convert-my-file"
                    }
                }
            })

            import_task = cloudconvert.Task.upload(file_name=input_pdf_path, task=job['tasks'][0])
            job = cloudconvert.Job.wait(id=job['id'])
            
            #Espera a que el estatus del job sea finalizado para proceder con la descarga del archivo convertido
            if job['status'] == 'finished':
                export_task_id = job['tasks'][0]['id']
                export_task = cloudconvert.Task.find(id=export_task_id)
                export_url = export_task['result']['files'][0]['url']
                response = requests.get(export_url)

                if response.status_code == 200:
                    with open(output_word_path, 'wb') as output_file:
                        output_file.write(response.content)
                    print("PDF converted to Word successfully!")
                    return True
                else:
                    print("Error downloading the converted file.")
            elif job['status'] == 'error' and job['error']['code'] == 'CREDITS_EXCEEDED':
                print(f"API Key {api_key} has run out of conversion credits.")
                continue  # Intenta con la siguiente api key registrada
            else:
                print(f"Conversion failed with API Key {api_key}.")
        except cloudconvert.exceptions.ClientError as e:
            print(f"Error with API Key {api_key}: {e}")
    print("Conversion failed with all API keys.")
    return False


                     